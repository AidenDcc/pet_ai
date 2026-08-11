# -*- coding: utf-8 -*-
"""
数心智能宠物健康平台 数据库表设计说明书 V1.0 生成脚本
- 参照《数心智能运动健康平台_数据库表设计说明书_V2.0》的章节结构与表格样式
- 每张表均包含：ID 主键（自增长）+ create_by / create_date / last_update_by / last_update_date / del_flag
- 输出：Word 文档（.docx）+ ER 图（.png）
"""
import os
import math
import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = BASE_DIR
OUT_DOCX = os.path.join(OUT_DIR, '数心智能宠物健康平台_数据库表设计说明书_V1.0.docx')
ER_PNG = os.path.join(OUT_DIR, 'er_diagram.png')

# ---------------------------------------------------------------
# 字体：匹配参考文档（Calibri + 微软雅黑）
# ---------------------------------------------------------------
LATIN = 'Calibri'
EAST = '微软雅黑'
ACCENT = RGBColor(0x47, 0x55, 0x69)      # 参考文档主色 #475569
HEADER_BG = '475569'
ZEBRA_BG = 'F2F5F9'
BORDER_COLOR = 'C9D3E0'
FONT_SMALL = 9

# ---------------------------------------------------------------
# 标准审计字段（每张表都需包含）
# ---------------------------------------------------------------
AUDIT_FIELDS = [
    ('create_by',        'BIGINT',   '-',   'Y', '0',                 '创建人ID'),
    ('create_date',      'DATETIME', '-',   'Y', 'CURRENT_TIMESTAMP', '创建时间'),
    ('last_update_by',   'BIGINT',   '-',   'Y', '0',                 '最后更新人ID'),
    ('last_update_date', 'DATETIME', '-',   'Y', 'CURRENT_TIMESTAMP', '最后更新时间'),
    ('del_flag',         'TINYINT',  '1',   'Y', '0',                 '删除标记：0-正常，1-已删除'),
]


def table(name, desc, fields, notes=None):
    """构造一张表：ID 主键 + 业务字段 + 审计字段"""
    rows = [('ID', 'BIGINT', '-', 'Y', 'AUTO_INCREMENT', '主键，自增长')]
    rows += fields
    rows += AUDIT_FIELDS
    return {'name': name, 'desc': desc, 'rows': rows, 'notes': notes or []}


# ---------------------------------------------------------------
# 3. 用户管理模块
# ---------------------------------------------------------------
T_USER = table('user', '宠物主账号表：平台 C 端宠物主（含宠物医生）登录账号',
    [
        ('account',          'VARCHAR',  '50',      'Y', '-',                 '登录账号（唯一）'),
        ('password',         'VARCHAR',  '100',     'Y', '-',                 '登录密码（BCrypt 加密存储）'),
        ('nickname',         'VARCHAR',  '100',     'N', '-',                 '用户昵称'),
        ('avatar_url',       'VARCHAR',  '500',     'N', '-',                 '头像URL'),
        ('phone',            'VARCHAR',  '20',      'N', '-',                 '手机号（加密存储）'),
        ('role',             'TINYINT',  '1',       'Y', '1',                 '账号角色：1-宠物主，2-宠物医生'),
        ('status',           'TINYINT',  '1',       'Y', '1',                 '账号状态：0-禁用，1-正常'),
        ('registered_date',  'DATETIME', '-',       'Y', 'CURRENT_TIMESTAMP', '注册时间'),
        ('last_active_date', 'DATETIME', '-',       'N', '-',                 '最后活跃时间'),
    ],
    ['account 建立唯一索引 uk_user_account。',
     '平台运营内部账号独立于本表，使用 sys_user 管理。'])

T_USER_SUB = table('user_subscription', '用户订阅关系表：记录用户与订阅套餐的绑定（含历史）',
    [
        ('user_id',      'BIGINT',   '-',   'Y', '-',                 '用户ID'),
        ('plan_id',      'BIGINT',   '-',   'Y', '-',                 '订阅套餐ID'),
        ('status',       'TINYINT',  '1',   'Y', '1',                 '订阅状态：0-已退订/已过期，1-生效中'),
        ('start_date',   'DATETIME', '-',   'Y', 'CURRENT_TIMESTAMP', '生效时间'),
        ('expire_date',  'DATETIME', '-',   'Y', '-',                 '到期时间'),
    ],
    ['user_id 与 plan_id 组合索引 idx_user_sub_plan，支撑按用户查询当前订阅。'])

T_PET = table('pet', '宠物档案表：宠物基本信息与健康档案',
    [
        ('owner_id',         'VARCHAR',  '50',      'Y', '-',           '宠物主用户ID'),
        ('pet_name',         'VARCHAR',  '100',     'Y', '-',           '宠物昵称'),
        ('species',          'TINYINT',  '1',       'Y', '-',           '物种：1-犬，2-猫'),
        ('breed',            'VARCHAR',  '100',     'N', '-',           '品种'),
        ('gender',           'TINYINT',  '1',       'Y', '-',           '性别：1-公，2-母'),
        ('birth_date',       'DATE',     '-',       'N', '-',           '出生日期'),
        ('weight',           'DECIMAL',  '6,2',     'N', '-',           '体重（kg）'),
        ('avatar_url',       'VARCHAR',  '500',     'N', '-',           '头像URL'),
        ('sterilized',       'TINYINT',  '1',       'Y', '0',           '是否绝育：0-否，1-是'),
        ('microchip',        'VARCHAR',  '50',      'N', '-',           '宠物芯片编号'),
        ('personality_tags', 'VARCHAR',  '200',     'N', '-',           '性格标签（英文逗号分隔，最多10个）'),
        ('device_id',        'BIGINT',   '-',       'N', '-',           '当前绑定设备ID（冗余，绑定关系见 user_device_bind）'),
    ],
    ['owner_id 建立索引 idx_pet_owner，device_id 建立索引 idx_pet_device。'])

T_VACCINE = table('pet_vaccine', '疫苗记录表：宠物疫苗接种历史',
    [
        ('pet_id',        'BIGINT',  '-',    'Y', '-', '宠物ID'),
        ('vaccine_name',  'VARCHAR', '100',  'Y', '-', '疫苗名称（如：犬八联疫苗、狂犬病疫苗）'),
        ('vaccine_date',  'DATE',    '-',    'Y', '-', '接种日期'),
        ('note',          'VARCHAR', '500',  'N', '-', '备注（如：第几针）'),
    ],
    ['pet_id 建立索引 idx_vaccine_pet。'])

T_DEWORM = table('pet_deworm', '驱虫记录表：宠物体内/体外驱虫历史',
    [
        ('pet_id',       'BIGINT',  '-',    'Y', '-', '宠物ID'),
        ('deworm_name',  'VARCHAR', '100',  'Y', '-', '驱虫名称（如：体内驱虫、体外驱虫）'),
        ('deworm_date',  'DATE',    '-',    'Y', '-', '驱虫日期'),
        ('note',         'VARCHAR', '500',  'N', '-', '备注'),
    ],
    ['pet_id 建立索引 idx_deworm_pet。'])

T_SYS_USER = table('sys_user', '系统用户表：平台运营/后台内部账号',
    [
        ('username',         'VARCHAR',  '50',      'Y', '-',                 '登录账号（唯一）'),
        ('password',         'VARCHAR',  '100',     'Y', '-',                 '登录密码（加密存储）'),
        ('real_name',        'VARCHAR',  '100',     'Y', '-',                 '姓名'),
        ('role_id',          'BIGINT',   '-',       'Y', '-',                 '角色ID'),
        ('phone',            'VARCHAR',  '20',      'N', '-',                 '手机号'),
        ('email',            'VARCHAR',  '100',     'N', '-',                 '邮箱'),
        ('status',           'TINYINT',  '1',       'Y', '1',                 '账号状态：0-禁用，1-正常'),
        ('last_login_date',  'DATETIME', '-',       'N', '-',                 '最后登录时间'),
    ],
    ['username 建立唯一索引 uk_sys_user_username。'])

T_SYS_ROLE = table('sys_role', '角色表：后台 RBAC 角色',
    [
        ('role_name',  'VARCHAR', '100',  'Y', '-', '角色名称'),
        ('role_code',  'VARCHAR', '50',   'Y', '-', '角色编码（唯一，如 admin/operator/auditor）'),
        ('sort',       'INT',     '11',   'Y', '0', '显示排序'),
        ('status',     'TINYINT', '1',    'Y', '1', '状态：0-禁用，1-正常'),
        ('remark',     'VARCHAR', '500',  'N', '-', '备注'),
    ],
    ['role_code 建立唯一索引 uk_sys_role_code。'])

T_SYS_MENU = table('sys_menu', '菜单表：后台菜单与按钮权限',
    [
        ('parent_id',  'BIGINT',   '-',    'N', '-',           '父菜单ID（一级菜单为空）'),
        ('menu_name',  'VARCHAR',  '100',  'Y', '-',           '菜单名称'),
        ('menu_type',  'TINYINT',  '1',    'Y', '-',           '类型：1-目录，2-菜单，3-按钮'),
        ('icon',       'VARCHAR',  '100',  'N', '-',           '图标'),
        ('path',       'VARCHAR',  '200',  'N', '-',           '路由路径'),
        ('perm',       'VARCHAR',  '200',  'N', '-',           '权限标识（如 admin:user:add）'),
        ('sort',       'INT',      '11',   'Y', '0',           '显示排序'),
        ('visible',    'TINYINT',  '1',    'Y', '1',           '是否显示：0-隐藏，1-显示'),
        ('status',     'TINYINT',  '1',    'Y', '1',           '状态：0-禁用，1-正常'),
    ],
    ['parent_id 建立索引 idx_sys_menu_parent，perm 建立索引 idx_sys_menu_perm。'])

T_SYS_ROLE_MENU = table('sys_role_menu', '角色菜单关联表：角色与菜单/权限多对多',
    [
        ('role_id',  'BIGINT', '-', 'Y', '-', '角色ID'),
        ('menu_id',  'BIGINT', '-', 'Y', '-', '菜单ID'),
    ],
    ['role_id 与 menu_id 组合唯一索引 uk_role_menu。'])

# ---------------------------------------------------------------
# 4. 设备管理模块
# ---------------------------------------------------------------
T_DEVICE = table('device', '设备表：Pet-S1 智能项圈设备',
    [
        ('sn',              'VARCHAR',  '50',      'Y', '-',                 '设备SN（唯一）'),
        ('imei',            'VARCHAR',  '50',      'Y', '-',                 'IMEI 设备唯一标识'),
        ('device_name',     'VARCHAR',  '100',     'Y', '-',                 '设备名称（Pet-S1 智能项圈）'),
        ('model',           'VARCHAR',  '50',      'Y', '-',                 '设备型号：Pet-S1'),
        ('status',          'TINYINT',  '1',       'Y', '1',                 '设备状态：0-离线，1-在线，2-低电量，3-未绑定'),
        ('battery',         'TINYINT',  '1',       'Y', '100',               '电量（0-100）'),
        ('signal',          'INT',      '11',      'N', '-',                 '信号强度（dBm，-100 ~ -30）'),
        ('firmware',        'VARCHAR',  '50',      'N', '-',                 '固件版本'),
        ('owner_id',        'BIGINT',   '-',       'N', '-',                 '绑定宠物主ID'),
        ('bound_pet_id',    'BIGINT',   '-',       'N', '-',                 '绑定宠物ID'),
        ('cur_lng',         'DECIMAL',  '10,6',    'N', '-',                 '最新定位经度'),
        ('cur_lat',         'DECIMAL',  '10,6',    'N', '-',                 '最新定位纬度'),
        ('activated_at',    'DATETIME', '-',       'N', '-',                 '激活时间'),
        ('last_sync_date',  'DATETIME', '-',       'N', '-',                 '最后同步时间'),
    ],
    ['sn 建立唯一索引 uk_device_sn，imei 建立唯一索引 uk_device_imei。',
     '设备绑定关系以 user_device_bind 为准，本表 owner_id/bound_pet_id 为冗余便于查询。'])

T_BIND = table('user_device_bind', '用户设备绑定表：用户与设备的绑定/解绑历史',
    [
        ('user_id',     'BIGINT',   '-',       'Y', '-',                 '用户ID'),
        ('device_id',   'BIGINT',   '-',       'Y', '-',                 '设备ID'),
        ('pet_id',      'BIGINT',   '-',       'N', '-',                 '绑定宠物ID'),
        ('status',      'TINYINT',  '1',       'Y', '1',                 '绑定状态：0-已解绑，1-绑定中'),
        ('bind_date',   'DATETIME', '-',       'Y', 'CURRENT_TIMESTAMP', '绑定时间'),
        ('unbind_date', 'DATETIME', '-',       'N', '-',                 '解绑时间'),
    ],
    ['user_id 与 device_id 组合索引 idx_bind_user_device，device_id 建立索引 idx_bind_device。'])

T_CONNECT_LOG = table('device_connect_log', '设备连接日志表：设备上下线记录',
    [
        ('device_id',     'BIGINT',   '-',   'Y', '-',                 '设备ID'),
        ('status',        'TINYINT',  '1',   'Y', '-',                 '连接状态：1-上线，0-下线'),
        ('connect_date',  'DATETIME', '-',   'Y', '-',                 '连接时间'),
    ],
    ['device_id 建立索引 idx_connect_log_device，配合 connect_date 支撑查询。'])

T_UPLOAD_LOG = table('device_upload_log', '体征上报记录表：项圈上报完整生命体征的记录',
    [
        ('device_id',        'BIGINT',   '-',       'Y', '-',                 '设备ID'),
        ('pet_id',           'BIGINT',   '-',       'Y', '-',                 '宠物ID'),
        ('report_date',      'DATETIME', '-',       'Y', 'CURRENT_TIMESTAMP', '上报时间'),
        ('source',           'TINYINT',  '1',       'Y', '1',                 '上报来源：1-自动上报，2-手动上传'),
        ('status',           'TINYINT',  '1',       'Y', '1',                 '上报结果：0-失败，1-成功'),
        ('temperature',      'DECIMAL',  '5,2',     'N', '-',                 '体温（°C）'),
        ('heart_rate',       'INT',      '11',      'N', '-',                 '心率（次/分）'),
        ('spo2',             'INT',      '11',      'N', '-',                 '血氧饱和度（%）'),
        ('respiratory_rate', 'INT',      '11',      'N', '-',                 '呼吸频率（次/分）'),
    ],
    ['device_id 与 report_date 组合索引 idx_upload_device_date。'])

T_COMMAND_LOG = table('device_command_log', '远程指令记录表：下发到设备的远程指令',
    [
        ('device_id',     'BIGINT',   '-',    'Y', '-',                 '设备ID'),
        ('user_id',       'BIGINT',   '-',    'Y', '-',                 '下发用户ID'),
        ('command_type',  'VARCHAR',  '20',   'Y', '-',                 '指令类型：find-蜂鸣，light-亮灯，feed-投喂，refresh-刷新定位'),
        ('result',        'TINYINT',  '1',    'Y', '1',                 '执行结果：0-失败，1-成功'),
        ('fail_reason',   'VARCHAR',  '500',  'N', '-',                 '失败原因（如设备离线）'),
        ('command_date',  'DATETIME', '-',    'Y', 'CURRENT_TIMESTAMP', '指令下发时间'),
    ],
    ['device_id 建立索引 idx_command_device。'])

# ---------------------------------------------------------------
# 5. 健康监测模块
# ---------------------------------------------------------------
T_HEALTH_METRIC = table('pet_health_metric', '生命体征采样表：项圈按周期上报的生命体征明细',
    [
        ('pet_id',           'BIGINT',   '-',      'Y', '-',           '宠物ID'),
        ('device_id',        'BIGINT',   '-',      'N', '-',           '采集设备ID（冗余）'),
        ('metric_date',      'DATETIME', '-',      'Y', '-',           '采样时间'),
        ('heart_rate',       'INT',      '11',     'Y', '-',           '心率（次/分）'),
        ('respiratory_rate', 'INT',      '11',     'Y', '-',           '呼吸频率（次/分）'),
        ('spo2',             'INT',      '11',     'Y', '-',           '血氧饱和度（%）'),
        ('temperature',      'DECIMAL',  '5,2',    'Y', '-',           '体温（°C）'),
        ('activity',         'INT',      '11',     'Y', '0',           '活动量（该区间步数）'),
        ('sleep_stage',      'TINYINT',  '1',      'Y', '0',           '睡眠阶段：0-清醒，1-浅睡，2-深睡'),
    ],
    ['高频数据表，按 pet_id + metric_date 建立索引 idx_health_metric_pet_date，建议按月分区并定期归档。'])

T_HEALTH_DAILY = table('pet_health_daily', '健康日汇总表：宠物每日健康数据汇总',
    [
        ('pet_id',         'BIGINT',   '-',     'Y', '-',  '宠物ID'),
        ('stat_date',      'DATE',     '-',     'Y', '-',  '统计日期'),
        ('steps',          'INT',      '11',    'Y', '0',  '当日总步数'),
        ('sleep_hours',    'DECIMAL',  '5,2',   'Y', '0',  '当日睡眠时长（小时）'),
        ('avg_heart_rate', 'INT',      '11',    'N', '-',  '当日平均心率（次/分）'),
    ],
    ['pet_id 与 stat_date 组合唯一索引 uk_health_daily_pet_date。'])

T_EXERCISE = table('pet_exercise', '运动监测表：宠物运动状态实时数据',
    [
        ('pet_id',       'BIGINT',   '-',     'Y', '-',           '宠物ID'),
        ('device_id',    'BIGINT',   '-',     'N', '-',           '采集设备ID'),
        ('metric_date',  'DATETIME', '-',     'Y', '-',           '采样时间'),
        ('step_freq',    'INT',      '11',    'Y', '0',           '步频（步/分）'),
        ('stride',       'INT',      '11',    'Y', '0',           '步幅（cm）'),
        ('speed',        'DECIMAL',  '5,2',   'Y', '0',           '速度（m/s）'),
        ('gait',         'TINYINT',  '1',     'Y', '1',           '步态：1-静止，2-行走，3-小跑，4-奔跑'),
    ],
    ['pet_id 与 metric_date 组合索引 idx_exercise_pet_date。'])

T_LOCATION_TRACK = table('pet_location_track', '定位轨迹表：宠物实时定位与历史轨迹',
    [
        ('pet_id',       'BIGINT',   '-',       'Y', '-',           '宠物ID'),
        ('device_id',    'BIGINT',   '-',       'Y', '-',           '设备ID'),
        ('lng',          'DECIMAL',  '10,6',    'Y', '-',           '经度'),
        ('lat',          'DECIMAL',  '10,6',    'Y', '-',           '纬度'),
        ('track_date',   'DATETIME', '-',       'Y', '-',           '定位时间'),
        ('address',      'VARCHAR',  '500',     'N', '-',           '定位物理地址（省市区，可选冗余）'),
    ],
    ['pet_id 与 track_date 组合索引 idx_track_pet_date。'])

T_FENCE = table('pet_fence', '电子围栏表：宠物安全区域（每只宠物可配置多个）',
    [
        ('pet_id',      'BIGINT',   '-',       'Y', '-',           '宠物ID'),
        ('fence_name',  'VARCHAR',  '100',     'Y', '-',           '围栏名称（如：小区、公园）'),
        ('center_lng',  'DECIMAL',  '10,6',    'Y', '-',           '中心点经度'),
        ('center_lat',  'DECIMAL',  '10,6',    'Y', '-',           '中心点纬度'),
        ('radius',      'INT',      '11',      'Y', '500',         '围栏半径（米）'),
        ('enabled',     'TINYINT',  '1',       'Y', '1',           '是否启用：0-关闭，1-启用'),
        ('address',     'VARCHAR',  '500',     'N', '-',           '中心点物理地址'),
    ],
    ['pet_id 建立索引 idx_fence_pet。'])

T_FENCE_ALERT = table('pet_fence_alert', '围栏告警记录表：宠物出入围栏触发告警',
    [
        ('pet_id',      'BIGINT',   '-',       'Y', '-',                 '宠物ID'),
        ('fence_id',    'BIGINT',   '-',       'Y', '-',                 '围栏ID'),
        ('device_id',   'BIGINT',   '-',       'Y', '-',                 '设备ID'),
        ('alert_type',  'TINYINT',  '1',       'Y', '1',                 '告警类型：1-出围栏，2-入围栏'),
        ('lng',         'DECIMAL',  '10,6',    'N', '-',                 '触发位置经度'),
        ('lat',         'DECIMAL',  '10,6',    'N', '-',                 '触发位置纬度'),
        ('alert_date',  'DATETIME', '-',       'Y', 'CURRENT_TIMESTAMP', '告警时间'),
        ('status',      'TINYINT',  '1',       'Y', '0',                 '处理状态：0-未处理，1-已处理'),
    ],
    ['pet_id 与 alert_date 组合索引 idx_fence_alert_pet_date。'])

# ---------------------------------------------------------------
# 6. 报告中心模块
# ---------------------------------------------------------------
T_REPORT = table('health_report', '健康报告表：周期健康评估报告（含 AI 结论与医生审核）',
    [
        ('pet_id',               'BIGINT',   '-',       'Y', '-',                 '宠物ID'),
        ('report_type',          'TINYINT',  '1',       'Y', '1',                 '报告周期：1-周报，2-月报'),
        ('period_start',         'DATE',     '-',       'Y', '-',                 '统计开始日期'),
        ('period_end',           'DATE',     '-',       'Y', '-',                 '统计结束日期'),
        ('score',                'INT',      '11',      'Y', '-',                 '健康综合评分（0-100）'),
        ('summary',              'VARCHAR',  '2000',    'N', '-',                 '报告摘要'),
        ('ai_conclusion',        'TEXT',     '-',       'N', '-',                 'AI 智能分析结论'),
        ('review_status',        'TINYINT',  '1',       'Y', '0',                 '审核状态：0-待审核，1-已通过，2-已驳回'),
        ('doctor_id',            'BIGINT',   '-',       'N', '-',                 '审核医生ID'),
        ('doctor_comment',       'VARCHAR',  '1000',    'N', '-',                 '医生审核意见'),
        ('avg_heart_rate',       'INT',      '11',      'N', '-',                 '平均心率'),
        ('max_heart_rate',       'INT',      '11',      'N', '-',                 '最高心率'),
        ('min_heart_rate',       'INT',      '11',      'N', '-',                 '最低心率'),
        ('avg_respiratory_rate', 'INT',      '11',      'N', '-',                 '平均呼吸频率'),
        ('avg_spo2',             'INT',      '11',      'N', '-',                 '平均血氧（%）'),
        ('avg_temperature',      'DECIMAL',  '5,2',     'N', '-',                 '平均体温（°C）'),
        ('total_activity',       'INT',      '11',      'N', '0',                 '周期内总活动步数'),
        ('sleep_duration',       'DECIMAL',  '5,2',     'N', '0',                 '周期内平均睡眠时长（小时）'),
        ('report_date',          'DATETIME', '-',       'Y', 'CURRENT_TIMESTAMP', '报告生成时间'),
    ],
    ['pet_id 建立索引 idx_report_pet，review_status 建立索引 idx_report_review（医生待审列表）。',
     '指标汇总（metricsSummary）拆分为列存储，异常项见 report_abnormal_item。'])

T_REPORT_ABNORMAL = table('report_abnormal_item', '报告异常项表：健康报告中的异常指标明细',
    [
        ('report_id',       'BIGINT',   '-',    'Y', '-',           '报告ID'),
        ('abnormal_key',    'VARCHAR',  '50',   'Y', '-',           '异常指标编码（hr/activity/sleep/spo2/temp）'),
        ('abnormal_label',  'VARCHAR',  '100',  'Y', '-',           '异常项名称（如：夜间心率）'),
        ('abnormal_value',  'VARCHAR',  '100',  'N', '-',           '异常数值/描述（如：偏高）'),
        ('level',           'TINYINT',  '1',    'Y', '1',           '异常等级：1-提示，2-警告，3-危险'),
        ('suggestion',      'VARCHAR',  '1000', 'N', '-',           '健康建议'),
    ],
    ['report_id 建立索引 idx_report_abnormal_report。'])

# ---------------------------------------------------------------
# 7. 在线问诊模块
# ---------------------------------------------------------------
T_VET = table('vet', '宠物医生表：入驻平台的宠物医生档案',
    [
        ('user_id',         'BIGINT',   '-',       'Y', '-',           '关联登录账号ID（user.id，role=2）'),
        ('doctor_name',     'VARCHAR',  '100',     'Y', '-',           '医生姓名'),
        ('hospital',        'VARCHAR',  '200',     'Y', '-',           '所属医院'),
        ('doctor_title',    'VARCHAR',  '100',     'N', '-',           '职称（如：主任医师）'),
        ('avatar_url',      'VARCHAR',  '500',     'N', '-',           '头像URL'),
        ('specialty',       'VARCHAR',  '200',     'N', '-',           '擅长领域'),
        ('consult_price',   'DECIMAL',  '8,2',     'Y', '0',           '问诊定价（元）'),
        ('phone',           'VARCHAR',  '20',      'N', '-',           '联系电话'),
        ('cert_status',     'TINYINT',  '1',       'Y', '0',           '认证状态：0-待审核，1-已通过，2-已驳回'),
        ('cert_img_url',    'VARCHAR',  '500',     'N', '-',           '执业证书图片URL'),
    ],
    ['user_id 建立唯一索引 uk_vet_user，cert_status 建立索引 idx_vet_cert。'])

T_CONSULT = table('consultation', '问诊记录表：宠物主将健康数据推送给医生的问诊',
    [
        ('pet_id',     'BIGINT',   '-',    'Y', '-',                 '宠物ID'),
        ('owner_id',   'BIGINT',   '-',    'Y', '-',                 '宠物主用户ID'),
        ('doctor_id',  'BIGINT',   '-',    'Y', '-',                 '医生ID（vet.id）'),
        ('status',     'TINYINT',  '1',    'Y', '1',                 '问诊状态：0-已关闭，1-进行中'),
        ('pushed_at',  'DATETIME', '-',    'Y', 'CURRENT_TIMESTAMP', '健康数据推送时间'),
        ('note',       'VARCHAR',  '1000', 'N', '-',                 '问诊描述（症状说明）'),
    ],
    ['doctor_id 与 status 组合索引 idx_consult_doctor_status，owner_id 建立索引 idx_consult_owner。',
     '同一宠物同一医生同一时间段的重复推送视为更新（唯一键 pet_id+doctor_id+status 有效时）。'])

T_CONSULT_MSG = table('consultation_message', '问诊消息表：问诊过程交流消息（扩展预留）',
    [
        ('consultation_id',  'BIGINT',   '-',    'Y', '-',                 '问诊记录ID'),
        ('sender_type',      'TINYINT',  '1',    'Y', '1',                 '发送方：1-宠物主，2-医生'),
        ('sender_id',        'BIGINT',   '-',    'Y', '-',                 '发送人ID（user.id）'),
        ('content',          'TEXT',     '-',    'Y', '-',                 '消息内容'),
        ('message_date',     'DATETIME', '-',    'Y', 'CURRENT_TIMESTAMP', '发送时间'),
    ],
    ['consultation_id 建立索引 idx_consult_msg_consult。'])

# ---------------------------------------------------------------
# 8. 订阅计费模块
# ---------------------------------------------------------------
T_PLAN = table('subscription_plan', '订阅套餐表：平台订阅套餐配置',
    [
        ('plan_name',       'VARCHAR',  '100',   'Y', '-',           '套餐名称'),
        ('plan_code',       'VARCHAR',  '50',    'Y', '-',           '套餐编码（唯一，如 basic/pro/premium）'),
        ('price',           'DECIMAL',  '10,2',  'Y', '-',           '价格（元/年）'),
        ('duration_months', 'INT',      '11',    'Y', '12',          '有效期（月）'),
        ('features',        'TEXT',     '-',     'N', '-',           '功能权益（多个以分号分隔）'),
        ('plan_color',      'VARCHAR',  '20',    'N', '-',           '标识色（前端展示）'),
        ('sort',            'INT',      '11',    'Y', '0',           '显示排序'),
        ('status',          'TINYINT',  '1',     'Y', '1',           '状态：0-下架，1-上架'),
    ],
    ['plan_code 建立唯一索引 uk_plan_code。'])

T_ORDER = table('order', '订阅订单表：订阅套餐购买/续费订单',
    [
        ('order_no',     'VARCHAR',  '50',      'Y', '-',                 '订单号（唯一）'),
        ('user_id',      'BIGINT',   '-',       'Y', '-',                 '下单用户ID'),
        ('pet_id',       'BIGINT',   '-',       'N', '-',                 '关联宠物ID（可空）'),
        ('plan_id',      'BIGINT',   '-',       'Y', '-',                 '套餐ID'),
        ('plan_name',    'VARCHAR',  '100',     'Y', '-',                 '套餐名称（下单快照）'),
        ('amount',       'DECIMAL',  '10,2',    'Y', '-',                 '订单金额（元）'),
        ('status',       'TINYINT',  '1',       'Y', '0',                 '订单状态：0-待支付，1-已支付，2-已过期，3-已退款'),
        ('pay_method',   'VARCHAR',  '50',      'N', '-',                 '支付方式（微信支付/支付宝等）'),
        ('pay_date',     'DATETIME', '-',       'N', '-',                 '支付时间'),
        ('expire_date',  'DATETIME', '-',       'N', '-',                 '订单有效期（超时未付自动过期）'),
    ],
    ['order_no 建立唯一索引 uk_order_no，user_id 建立索引 idx_order_user。'])

# ---------------------------------------------------------------
# 9. 宠物社区模块
# ---------------------------------------------------------------
T_POST = table('community_post', '社区帖子表：宠物社区动态',
    [
        ('author_id',      'BIGINT',   '-',       'Y', '-',           '作者用户ID'),
        ('pet_id',         'BIGINT',   '-',       'N', '-',           '关联宠物ID（可空）'),
        ('pet_name',       'VARCHAR',  '100',     'N', '-',           '宠物昵称（发布时快照）'),
        ('content',        'VARCHAR',  '2000',    'N', '-',           '帖子文案'),
        ('images',         'TEXT',     '-',       'N', '-',           '图片列表（JSON 数组）'),
        ('view_count',     'INT',      '11',      'Y', '0',           '浏览量'),
        ('like_count',     'INT',      '11',      'Y', '0',           '点赞量（冗余）'),
        ('comment_count',  'INT',      '11',      'Y', '0',           '评论量（冗余）'),
        ('status',         'TINYINT',  '1',       'Y', '1',           '状态：0-已下架/封禁，1-正常'),
    ],
    ['author_id 建立索引 idx_post_author，create_date 建立索引 idx_post_create。'])

T_COMMENT = table('community_comment', '社区评论表：帖子评论与回复',
    [
        ('post_id',    'BIGINT',   '-',    'Y', '-',           '帖子ID'),
        ('author_id',  'BIGINT',   '-',    'Y', '-',           '评论人用户ID'),
        ('parent_id',  'BIGINT',   '-',    'N', '-',           '父评论ID（回复评论时为空）'),
        ('content',    'VARCHAR',  '1000', 'Y', '-',           '评论内容'),
    ],
    ['post_id 建立索引 idx_comment_post，parent_id 建立索引 idx_comment_parent。'])

T_LIKE = table('community_like', '社区点赞表：帖子点赞记录',
    [
        ('post_id',  'BIGINT', '-', 'Y', '-', '帖子ID'),
        ('user_id',  'BIGINT', '-', 'Y', '-', '点赞用户ID'),
    ],
    ['post_id 与 user_id 组合唯一索引 uk_like_post_user，防止重复点赞。'])

T_FOLLOW = table('community_follow', '关注关系表：用户关注其他萌宠账号',
    [
        ('follower_id',  'BIGINT', '-', 'Y', '-', '关注者用户ID'),
        ('target_id',    'BIGINT', '-', 'Y', '-', '被关注者用户ID'),
    ],
    ['follower_id 与 target_id 组合唯一索引 uk_follow_follower_target，防止重复关注。'])

# ---------------------------------------------------------------
# 10. AI 助手模块
# ---------------------------------------------------------------
T_ASST_SESSION = table('assistant_session', 'AI 会话表：用户与 AI 助手的会话',
    [
        ('user_id',        'BIGINT',  '-',    'Y', '-',           '用户ID'),
        ('pet_id',         'BIGINT',  '-',    'Y', '-',           '关联宠物ID'),
        ('session_title',  'VARCHAR', '200',  'N', '-',           '会话标题（首轮问题摘要）'),
        ('status',         'TINYINT', '1',    'Y', '1',           '会话状态：0-已关闭，1-进行中'),
    ],
    ['user_id 建立索引 idx_asst_session_user。'])

T_ASST_MSG = table('assistant_message', 'AI 消息表：会话中的用户提问与 AI 回复',
    [
        ('session_id',   'BIGINT',   '-',    'Y', '-',                 '会话ID'),
        ('sender_type',  'TINYINT',  '1',    'Y', '1',                 '消息方向：1-用户提问，2-AI回复'),
        ('content',      'TEXT',     '-',    'Y', '-',                 '消息内容'),
        ('intent',       'VARCHAR',  '50',   'N', '-',                 '识别意图（如 find/light/location/health/report）'),
        ('message_date', 'DATETIME', '-',    'Y', 'CURRENT_TIMESTAMP', '发送时间'),
    ],
    ['session_id 建立索引 idx_asst_msg_session。'])

# ---------------------------------------------------------------
# 11. 系统管理模块
# ---------------------------------------------------------------
T_DICT_TYPE = table('dict_type', '字典类型表：数据字典类型',
    [
        ('dict_name',  'VARCHAR', '100',  'Y', '-', '字典名称'),
        ('dict_type',  'VARCHAR', '50',   'Y', '-', '字典类型编码（唯一）'),
        ('remark',     'VARCHAR', '500',  'N', '-', '备注'),
    ],
    ['dict_type 建立唯一索引 uk_dict_type_code。'])

T_DICT_ITEM = table('dict_item', '字典项表：数据字典明细',
    [
        ('type_id',     'BIGINT',   '-',    'Y', '-',           '字典类型ID'),
        ('item_label',  'VARCHAR',  '100',  'Y', '-',           '字典项名称（中文）'),
        ('item_value',  'VARCHAR',  '100',  'Y', '-',           '字典项值'),
        ('sort',        'INT',      '11',   'Y', '0',           '显示排序'),
        ('status',      'TINYINT',  '1',    'Y', '1',           '状态：0-禁用，1-正常'),
    ],
    ['type_id 建立索引 idx_dict_item_type。'])

T_LOGIN_LOG = table('login_log', '登录日志表：系统用户登录记录',
    [
        ('username',   'VARCHAR',  '50',    'Y', '-',                 '登录账号'),
        ('ip',         'VARCHAR',  '50',    'N', '-',                 '登录IP'),
        ('location',   'VARCHAR',  '200',   'N', '-',                 '登录地点'),
        ('browser',    'VARCHAR',  '100',   'N', '-',                 '浏览器'),
        ('os',         'VARCHAR',  '100',   'N', '-',                 '操作系统'),
        ('status',     'TINYINT',  '1',     'Y', '1',                 '登录结果：0-失败，1-成功'),
        ('message',    'VARCHAR',  '500',   'N', '-',                 '结果描述（密码错误/账号不存在等）'),
        ('login_date', 'DATETIME', '-',     'Y', 'CURRENT_TIMESTAMP', '登录时间'),
    ],
    ['login_date 建立索引 idx_login_log_date。'])

T_TERMINAL = table('terminal', '客户端终端表：平台各端（APP/H5/小程序/PC）配置',
    [
        ('terminal_name',   'VARCHAR',  '100',  'Y', '-',           '终端名称'),
        ('terminal_code',   'VARCHAR',  '50',   'Y', '-',           '终端编码（唯一）'),
        ('terminal_type',   'TINYINT',  '1',    'Y', '-',           '终端类型：1-APP，2-H5，3-小程序，4-PC'),
        ('latest_version',  'VARCHAR',  '50',   'N', '-',           '最新版本号'),
        ('download_url',    'VARCHAR',  '500',  'N', '-',           '下载地址'),
        ('status',          'TINYINT',  '1',    'Y', '1',           '状态：0-禁用，1-启用'),
        ('remark',          'VARCHAR',  '500',  'N', '-',           '备注'),
    ],
    ['terminal_code 建立唯一索引 uk_terminal_code。'])

T_I18N = table('i18n_entry', '国际化词条表：运营端可在线配置的中英文词条',
    [
        ('entry_key',    'VARCHAR',  '2000', 'Y', '-',           '词条Key（唯一）'),
        ('zh_cn',        'VARCHAR',  '2000', 'Y', '-',           '中文文案'),
        ('en_us',        'VARCHAR',  '2000', 'Y', '-',           '英文文案'),
        ('is_override',  'TINYINT',  '1',    'Y', '0',           '是否运营端覆盖基础词条：0-基础，1-覆盖'),
    ],
    ['entry_key 建立唯一索引 uk_i18n_key。'])

# ---------------------------------------------------------------
# 章节结构
# ---------------------------------------------------------------
SECTIONS = [
    ('3. 用户管理模块', [
        ('3.1 用户表（user）', T_USER),
        ('3.2 用户订阅关系表（user_subscription）', T_USER_SUB),
        ('3.3 宠物档案表（pet）', T_PET),
        ('3.4 疫苗记录表（pet_vaccine）', T_VACCINE),
        ('3.5 驱虫记录表（pet_deworm）', T_DEWORM),
        ('3.6 系统用户表（sys_user）', T_SYS_USER),
        ('3.7 角色表（sys_role）', T_SYS_ROLE),
        ('3.8 菜单表（sys_menu）', T_SYS_MENU),
        ('3.9 角色菜单关联表（sys_role_menu）', T_SYS_ROLE_MENU),
    ]),
    ('4. 设备管理模块', [
        ('4.1 设备表（device）', T_DEVICE),
        ('4.2 用户设备绑定表（user_device_bind）', T_BIND),
        ('4.3 设备连接日志表（device_connect_log）', T_CONNECT_LOG),
        ('4.4 体征上报记录表（device_upload_log）', T_UPLOAD_LOG),
        ('4.5 远程指令记录表（device_command_log）', T_COMMAND_LOG),
    ]),
    ('5. 健康监测模块', [
        ('5.1 生命体征采样表（pet_health_metric）', T_HEALTH_METRIC),
        ('5.2 健康日汇总表（pet_health_daily）', T_HEALTH_DAILY),
        ('5.3 运动监测表（pet_exercise）', T_EXERCISE),
        ('5.4 定位轨迹表（pet_location_track）', T_LOCATION_TRACK),
        ('5.5 电子围栏表（pet_fence）', T_FENCE),
        ('5.6 围栏告警记录表（pet_fence_alert）', T_FENCE_ALERT),
    ]),
    ('6. 报告中心模块', [
        ('6.1 健康报告表（health_report）', T_REPORT),
        ('6.2 报告异常项表（report_abnormal_item）', T_REPORT_ABNORMAL),
    ]),
    ('7. 在线问诊模块', [
        ('7.1 宠物医生表（vet）', T_VET),
        ('7.2 问诊记录表（consultation）', T_CONSULT),
        ('7.3 问诊消息表（consultation_message）', T_CONSULT_MSG),
    ]),
    ('8. 订阅计费模块', [
        ('8.1 订阅套餐表（subscription_plan）', T_PLAN),
        ('8.2 订阅订单表（order）', T_ORDER),
    ]),
    ('9. 宠物社区模块', [
        ('9.1 社区帖子表（community_post）', T_POST),
        ('9.2 社区评论表（community_comment）', T_COMMENT),
        ('9.3 社区点赞表（community_like）', T_LIKE),
        ('9.4 关注关系表（community_follow）', T_FOLLOW),
    ]),
    ('10. AI 助手模块', [
        ('10.1 AI 会话表（assistant_session）', T_ASST_SESSION),
        ('10.2 AI 消息表（assistant_message）', T_ASST_MSG),
    ]),
    ('11. 系统管理模块', [
        ('11.1 字典类型表（dict_type）', T_DICT_TYPE),
        ('11.2 字典项表（dict_item）', T_DICT_ITEM),
        ('11.3 登录日志表（login_log）', T_LOGIN_LOG),
        ('11.4 客户端终端表（terminal）', T_TERMINAL),
        ('11.5 国际化词条表（i18n_entry）', T_I18N),
    ]),
]

# ---------------------------------------------------------------
# 关系汇总
# ---------------------------------------------------------------
RELATIONS = [
    ('1',  'user',              'user_subscription',     'user_id',  '1:N', '一个用户可有多个订阅记录（当前+历史）'),
    ('2',  'user',              'pet',                   'owner_id', '1:N', '一个用户可养多只宠物'),
    ('3',  'pet',               'pet_vaccine',           'pet_id',   '1:N', '一只宠物可有多条疫苗记录'),
    ('4',  'pet',               'pet_deworm',            'pet_id',   '1:N', '一只宠物可有多条驱虫记录'),
    ('5',  'pet',               'pet_health_metric',     'pet_id',   '1:N', '一只宠物有多条生命体征采样'),
    ('6',  'pet',               'pet_health_daily',      'pet_id',   '1:N', '一只宠物有多天健康汇总'),
    ('7',  'pet',               'pet_exercise',          'pet_id',   '1:N', '一只宠物有多条运动数据'),
    ('8',  'pet',               'pet_location_track',    'pet_id',   '1:N', '一只宠物有多条定位轨迹'),
    ('9',  'pet',               'pet_fence',             'pet_id',   '1:N', '一只宠物可有多个电子围栏'),
    ('10', 'pet_fence',         'pet_fence_alert',       'fence_id', '1:N', '一个围栏可产生多条告警'),
    ('11', 'pet',               'health_report',         'pet_id',   '1:N', '一只宠物可有多份健康报告'),
    ('12', 'health_report',     'report_abnormal_item',  'report_id','1:N', '一份报告包含多条异常项'),
    ('13', 'user',              'user_device_bind',      'user_id',  '1:N', '一个用户可绑定多台设备'),
    ('14', 'device',            'user_device_bind',      'device_id','1:N', '一台设备可被多次绑定（含历史）'),
    ('15', 'pet',               'user_device_bind',      'pet_id',   '1:N', '一只宠物对应一条绑定记录'),
    ('16', 'device',            'device_connect_log',    'device_id','1:N', '一台设备有多条连接日志'),
    ('17', 'device',            'device_upload_log',     'device_id','1:N', '一台设备有多条体征上报'),
    ('18', 'device',            'device_command_log',    'device_id','1:N', '一台设备有多条远程指令'),
    ('19', 'vet',               'consultation',          'doctor_id','1:N', '一个医生可接多个问诊'),
    ('20', 'user',              'consultation',          'owner_id', '1:N', '一个用户可发起多次问诊'),
    ('21', 'pet',               'consultation',          'pet_id',   '1:N', '一只宠物可参与多次问诊'),
    ('22', 'consultation',      'consultation_message',  'consultation_id', '1:N', '一次问诊包含多条消息'),
    ('23', 'subscription_plan', 'user_subscription',     'plan_id',  '1:N', '一个套餐可被多个用户订阅'),
    ('24', 'subscription_plan', 'order',                 'plan_id',  '1:N', '一个套餐对应多笔订单'),
    ('25', 'user',              'order',                 'user_id',  '1:N', '一个用户可有多笔订单'),
    ('26', 'user',              'community_post',        'author_id','1:N', '一个用户可发布多篇帖子'),
    ('27', 'community_post',    'community_comment',     'post_id',  '1:N', '一篇帖子可有多条评论'),
    ('28', 'community_post',    'community_like',        'post_id',  '1:N', '一篇帖子可被多次点赞'),
    ('29', 'user',              'community_like',        'user_id',  '1:N', '一个用户可点赞多篇帖子'),
    ('30', 'user',              'community_follow',      'follower_id', '1:N', '一个用户可关注多个账号'),
    ('31', 'user',              'assistant_session',     'user_id',  '1:N', '一个用户可开启多个 AI 会话'),
    ('32', 'assistant_session', 'assistant_message',     'session_id','1:N', '一个会话包含多条消息'),
    ('33', 'sys_role',          'sys_user',              'role_id',  '1:N', '一个角色可分配给多个系统用户'),
    ('34', 'sys_role',          'sys_role_menu',         'role_id',  '1:N', '一个角色对应多条菜单权限'),
    ('35', 'sys_menu',          'sys_role_menu',         'menu_id',  '1:N', '一个菜单可分配给多个角色'),
    ('36', 'dict_type',         'dict_item',             'type_id',  '1:N', '一个字典类型包含多个字典项'),
]

# ---------------------------------------------------------------
# docx 渲染辅助
# ---------------------------------------------------------------
def set_run_font(run, size=10.5, bold=False, color=None, east=EAST, latin=LATIN):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), east)
    if color:
        run.font.color.rgb = color


def para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    return p


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell_border(cell, color=BORDER_COLOR):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def fill_cell(cell, text, size=FONT_SMALL, bold=False, color=None, align='center'):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def build_table(doc, header, rows, widths, left_cols=()):
    """统一表格构建：先建全部行，再统一设置列宽与行高。

    header: 表头列表
    rows:   数据行列表
    widths: 各列宽（Cm 对象）列表
    left_cols: 需要左对齐的列下标集合
    """
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        fill_cell(hdr[i], h, size=FONT_SMALL, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr[i], HEADER_BG)
        set_cell_border(hdr[i])

    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        zebra = ri % 2 == 1
        for ci, val in enumerate(row):
            align = 'left' if ci in left_cols else 'center'
            fill_cell(cells[ci], str(val), size=FONT_SMALL, align=align)
            if zebra:
                set_cell_bg(cells[ci], ZEBRA_BG)
            set_cell_border(cells[ci])

    # 全部行建立后统一设列宽
    for idx, w in enumerate(widths):
        for r in range(len(tbl.rows)):
            tbl.cell(r, idx).width = w

    # 行高控制
    for r in tbl.rows:
        trPr = r._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '300')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
    return tbl


def add_field_table(doc, rows, header=None):
    """7 列字段表：序号|字段名|数据类型|长度|必填|默认值|说明"""
    header = header or ['序号', '字段名', '数据类型', '长度', '必填', '默认值', '说明']
    widths = [Cm(0.9), Cm(3.0), Cm(2.3), Cm(1.3), Cm(1.0), Cm(2.6), Cm(5.9)]
    return build_table(doc, header, rows, widths, left_cols=(6,))


def add_field_table_4col(doc, rows):
    """5 列标准字段表：字段名|数据类型|必填|默认值|说明 (用于设计规范)"""
    header = ['字段名', '数据类型', '必填', '默认值', '说明']
    widths = [Cm(4.2), Cm(3.0), Cm(1.4), Cm(4.6), Cm(3.8)]
    return build_table(doc, header, rows, widths)


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 11), bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    if level == 1:
        p.paragraph_format.keep_with_next = True
        # 底部边框线
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '3')
        bottom.set(qn('w:color'), HEADER_BG)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = '右键更新域以生成目录（或打开文档后按 F9）。'
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)
    set_run_font(run, size=10.5, color=RGBColor(0x88, 0x88, 0x88))


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('数心智能宠物健康平台 · 数据库表设计说明书  |  第 ')
    set_run_font(run, size=8.5, color=RGBColor(0x8A, 0x8A, 0x8A))
    # PAGE field
    r = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2)
    set_run_font(r, size=8.5, color=RGBColor(0x8A, 0x8A, 0x8A))
    run2 = p.add_run(' 页')
    set_run_font(run2, size=8.5, color=RGBColor(0x8A, 0x8A, 0x8A))


# ---------------------------------------------------------------
# ER 图
# ---------------------------------------------------------------
def _er_font():
    candidates = [
        '/System/Library/AssetsV2/com_apple_MobileAsset_Font8/5feac9245cca79adaf638ded7a4994b1ddb33ca0.asset/AssetData/Hei.ttf',
        '/System/Library/Fonts/STHeiti Medium.ttc',
        '/System/Library/Fonts/STHeiti Light.ttc',
    ]
    for c in candidates:
        if os.path.exists(c):
            return fm.FontProperties(fname=c)
    return fm.FontProperties()  # 兜底：系统默认

ER_FONT = _er_font()

# box: name -> (cx, cy, w, h, key_fields)
ER_BOXES = {
    'user':               (2.0, 9.6, 3.1, 0.62, ['user 用户表', 'ID · account']),
    'pet':                (2.0, 7.9, 3.1, 0.62, ['pet 宠物档案表', 'ID · owner_id']),
    'pet_vaccine':        (5.2, 8.6, 3.1, 0.55, ['pet_vaccine 疫苗记录', 'ID · pet_id']),
    'pet_deworm':         (5.2, 7.9, 3.1, 0.55, ['pet_deworm 驱虫记录', 'ID · pet_id']),
    'pet_health_metric':  (5.2, 7.2, 3.1, 0.55, ['pet_health_metric 体征', 'ID · pet_id']),
    'health_report':      (5.2, 6.5, 3.1, 0.55, ['health_report 健康报告', 'ID · pet_id']),
    'device':             (2.0, 6.6, 3.1, 0.62, ['device 设备表', 'ID · sn · imei']),
    'user_device_bind':   (5.2, 5.8, 3.1, 0.55, ['user_device_bind 绑定', 'ID · device_id']),
    'user_subscription':  (5.2, 9.4, 3.1, 0.55, ['user_subscription 订阅', 'ID · user_id']),
    'subscription_plan':  (8.4, 9.4, 3.1, 0.55, ['subscription_plan 套餐', 'ID · plan_code']),
    'order':              (8.4, 8.6, 3.1, 0.55, ['order 订阅订单', 'ID · order_no']),
    'vet':                (8.4, 7.8, 3.1, 0.55, ['vet 宠物医生', 'ID · user_id']),
    'consultation':       (8.4, 7.0, 3.1, 0.55, ['consultation 问诊', 'ID · pet_id']),
    'community_post':     (8.4, 6.2, 3.1, 0.55, ['community_post 帖子', 'ID · author_id']),
    'community_comment':  (8.4, 5.4, 3.1, 0.55, ['community_comment 评论', 'ID · post_id']),
    'sys_user':           (2.0, 5.4, 3.1, 0.62, ['sys_user 系统用户', 'ID · username']),
    'sys_role':           (2.0, 4.4, 3.1, 0.55, ['sys_role 角色', 'ID · role_code']),
    'sys_menu':           (2.0, 3.6, 3.1, 0.55, ['sys_menu 菜单', 'ID · perm']),
    'sys_role_menu':      (4.6, 4.0, 3.1, 0.55, ['sys_role_menu 角色菜单', 'ID · role_id']),
}

# edge: (from, to, fk label)
ER_EDGES = [
    ('user', 'pet', 'owner_id'),
    ('pet', 'pet_vaccine', 'pet_id'),
    ('pet', 'pet_deworm', 'pet_id'),
    ('pet', 'pet_health_metric', 'pet_id'),
    ('pet', 'health_report', 'pet_id'),
    ('user', 'user_device_bind', 'user_id'),
    ('device', 'user_device_bind', 'device_id'),
    ('user', 'user_subscription', 'user_id'),
    ('subscription_plan', 'user_subscription', 'plan_id'),
    ('subscription_plan', 'order', 'plan_id'),
    ('user', 'order', 'user_id'),
    ('vet', 'consultation', 'doctor_id'),
    ('pet', 'consultation', 'pet_id'),
    ('user', 'community_post', 'author_id'),
    ('community_post', 'community_comment', 'post_id'),
    ('sys_role', 'sys_user', 'role_id'),
    ('sys_menu', 'sys_role_menu', 'menu_id'),
    ('sys_role', 'sys_role_menu', 'role_id'),
    ('pet', 'device', 'bound_pet_id'),
]


def box_edges(cx, cy, w, h):
    return cx - w / 2, cx + w / 2, cy - h / 2, cy + h / 2


def draw_box(ax, cx, cy, w, h, lines):
    box = FancyBboxPatch((cx - w / 2, cy - h / 2), w, h,
                         boxstyle='round,pad=0.03', linewidth=1.2,
                         edgecolor='#475569', facecolor='#EFF4FA', zorder=3)
    ax.add_patch(box)
    ax.text(cx, cy + h * 0.16, lines[0], ha='center', va='center',
            fontsize=10.5, fontproperties=ER_FONT, color='#1F2937', zorder=4)
    if len(lines) > 1:
        ax.text(cx, cy - h * 0.28, lines[1], ha='center', va='center',
                fontsize=8.2, fontproperties=ER_FONT, color='#64748B', zorder=4)


def draw_edge(ax, a, b, label):
    x1, y1 = ER_BOXES[a][0], ER_BOXES[a][1]
    x2, y2 = ER_BOXES[b][0], ER_BOXES[b][1]
    w1, h1 = ER_BOXES[a][2], ER_BOXES[a][3]
    w2, h2 = ER_BOXES[b][2], ER_BOXES[b][3]
    dx, dy = x2 - x1, y2 - y1
    L = math.hypot(dx, dy) or 1
    ux, uy = dx / L, dy / L
    # 从源盒边界到目标盒边界
    sx = x1 + ux * (w1 / 2 if abs(ux) > 0.02 else 0) + (uy * h1 / 2 if abs(uy) > 0.02 else 0)
    sy = y1 + uy * (h1 / 2 if abs(uy) > 0.02 else 0) + (ux * w1 / 2 if abs(ux) > 0.02 else 0)
    tx = x2 - ux * (w2 / 2 if abs(ux) > 0.02 else 0) - (uy * h2 / 2 if abs(uy) > 0.02 else 0)
    ty = y2 - uy * (h2 / 2 if abs(uy) > 0.02 else 0) - (ux * w2 / 2 if abs(ux) > 0.02 else 0)
    arrow = FancyArrowPatch((sx, sy), (tx, ty), arrowstyle='-|>', mutation_scale=12,
                            linewidth=1.0, color='#94A3B8', zorder=2, shrinkA=0, shrinkB=0)
    ax.add_patch(arrow)
    mx, my = (sx + tx) / 2, (sy + ty) / 2
    ax.text(mx + 0.06, my + 0.06, label, fontsize=7.5, fontproperties=ER_FONT,
            color='#475569', ha='center', va='center', zorder=5)


def render_er():
    fig, ax = plt.subplots(figsize=(12.5, 9.5), dpi=200)
    ax.set_xlim(0, 10.6)
    ax.set_ylim(2.6, 10.4)
    ax.axis('off')
    for name, (cx, cy, w, h, lines) in ER_BOXES.items():
        draw_box(ax, cx, cy, w, h, lines)
    for a, b, lab in ER_EDGES:
        draw_edge(ax, a, b, lab)
    ax.set_title('数心智能宠物健康平台 · 核心表 ER 图（示意）', fontproperties=ER_FONT,
                 fontsize=14, color='#1F2937', pad=14)
    fig.tight_layout()
    fig.savefig(ER_PNG, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return ER_PNG


# ---------------------------------------------------------------
# 文档生成
# ---------------------------------------------------------------
def build_doc():
    doc = Document()
    # 页面设置 A4，边距 2cm
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    add_page_number_footer(sec)

    # 默认样式
    normal = doc.styles['Normal']
    normal.font.name = LATIN
    normal.font.size = Pt(10.5)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), LATIN)
    rFonts.set(qn('w:hAnsi'), LATIN)
    rFonts.set(qn('w:eastAsia'), EAST)

    # ========== 封面 ==========
    for _ in range(5):
        para(doc, '', space_after=12)
    para(doc, '数心智能宠物健康平台', size=30, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, '数据库表设计说明书', size=40, bold=True, color=RGBColor(0x11, 0x1B, 0x28),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    para(doc, 'V1.0  |  2026-08-11', size=14, color=RGBColor(0x64, 0x74, 0x8B),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, '产品团队  |  技术架构组', size=12, color=RGBColor(0x64, 0x74, 0x8B),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # ========== 目录 ==========
    heading(doc, '目录', level=1)
    add_toc(doc)
    doc.add_page_break()

    # ========== 1. 引言 ==========
    heading(doc, '1. 引言', level=1)
    heading(doc, '1.1 目的', level=2)
    para(doc, '本文档是数心智能宠物健康平台的数据库表设计说明书，基于宠物主端（C端）、宠物医生端与平台运营端（B端）软件需求规格说明书编制。'
              '本文档详细定义了平台所有业务模块对应的数据库表结构，包括字段定义、数据类型、约束条件以及表间关系，'
              '作为后端开发团队进行数据库建模和开发的唯一权威依据。')
    heading(doc, '1.2 范围', level=2)
    para(doc, '本文档涵盖数心智能宠物健康平台全部业务模块的数据库设计，包括：用户管理模块、设备管理模块、健康监测模块、报告中心模块、'
              '在线问诊模块、订阅计费模块、宠物社区模块、AI 助手模块和系统管理模块。数据库基于 MySQL 8.0 设计，'
              '存储引擎采用 InnoDB，字符集使用 utf8mb4。')
    heading(doc, '1.3 术语和缩略语', level=2)
    terms = [
        ('Pet-S1 智能项圈', 'Pet-S1', '宠物智能穿戴设备，采集生命体征、定位与运动数据'),
        ('生命体征', 'Vital Signs', '心率、呼吸、血氧、体温等健康指标'),
        ('电子围栏', 'Geofence', '以指定位置为中心、按半径定义的虚拟安全区域'),
        ('健康报告', 'Health Report', '基于周期体征数据自动生成的健康评估报告'),
        ('健康评分', 'Health Score', '0-100 的综合健康评估分数'),
        ('在线问诊', 'Online Consultation', '宠物主将宠物健康数据推送给宠物医生进行咨询'),
        ('AI 智能分析', 'AI Analysis', '基于 AI 对健康数据的深度分析与疾病预警'),
        ('RBAC', 'RBAC', '基于角色的访问控制'),
    ]
    build_table(doc, ['术语', '英文', '定义'], terms,
                [Cm(4.0), Cm(3.5), Cm(9.5)], left_cols=(2,))

    # ========== 2. 设计规范 ==========
    doc.add_page_break()
    heading(doc, '2. 设计规范', level=1)
    heading(doc, '2.1 命名规范', level=2)
    for s in [
        '表名使用小写下划线命名法，如：pet_fence、health_report',
        '字段名使用小写下划线命名法，如：create_date、del_flag',
        '主键统一命名为 ID，自增长（BIGINT AUTO_INCREMENT）',
        '外键字段名格式：关联表名_id，如：user_id、pet_id',
        '索引命名：idx_表名_字段名；唯一索引：uk_表名_字段名',
        '字符集 utf8mb4、存储引擎 InnoDB',
    ]:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(s)
        set_run_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(3)
    heading(doc, '2.2 字段规范', level=2)
    para(doc, '所有表必须包含以下标准字段：')
    add_field_table_4col(doc, [
        ('create_by', 'BIGINT', 'Y', '0', '创建人ID'),
        ('create_date', 'DATETIME', 'Y', 'CURRENT_TIMESTAMP', '创建时间'),
        ('last_update_by', 'BIGINT', 'Y', '0', '最后更新人ID'),
        ('last_update_date', 'DATETIME', 'Y', 'CURRENT_TIMESTAMP', '最后更新时间'),
        ('del_flag', 'TINYINT', 'Y', '0', '删除标记：0-正常，1-已删除'),
    ])
    para(doc, '', space_after=2)
    para(doc, '状态字段统一使用 TINYINT 类型：0=禁用/删除，1=启用/正常；涉及多状态时按表中说明取值。',
         size=10, color=RGBColor(0x64, 0x74, 0x8B))

    # ========== 3-11 业务模块表结构 ==========
    for sec_title, tables in SECTIONS:
        doc.add_page_break()
        heading(doc, sec_title, level=1)
        for (sub_title, tdata) in tables:
            heading(doc, sub_title, level=2)
            para(doc, tdata['desc'], size=10, color=RGBColor(0x64, 0x74, 0x8B), space_after=4)
            # 字段序号
            rows = [(str(i + 1), f[0], f[1], f[2], f[3], f[4], f[5]) for i, f in enumerate(tdata['rows'])]
            add_field_table(doc, rows)
            para(doc, '', space_after=2)
            for n in tdata['notes']:
                p = doc.add_paragraph()
                run = p.add_run('说明：' + n)
                set_run_font(run, size=9, color=RGBColor(0x8A, 0x8A, 0x8A))
                p.paragraph_format.space_after = Pt(2)

    # ========== 12. ER 图 ==========
    doc.add_page_break()
    heading(doc, '12. 表关系 ER 图', level=1)
    para(doc, '下图为核心业务表之间的关联关系示意（完整关系见第 13 章「关系汇总表」）：')
    er_path = render_er()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(er_path, width=Inches(6.6))

    # ========== 13. 关系汇总表 ==========
    heading(doc, '13. 关系汇总表', level=1)
    para(doc, '全部表间关系如下：')
    rel_rows = [(r[0], r[1], r[2], r[3], r[4], r[5]) for r in RELATIONS]
    build_table(doc, ['序号', '主表', '关联表', '关联字段', '关系类型', '说明'], rel_rows,
                [Cm(1.1), Cm(3.4), Cm(3.4), Cm(3.0), Cm(1.8), Cm(4.3)], left_cols=(5,))

    doc.save(OUT_DOCX)
    print('已生成：', OUT_DOCX)
    print('已生成 ER 图：', ER_PNG)


if __name__ == '__main__':
    build_doc()
