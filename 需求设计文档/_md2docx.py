#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 PRD Markdown 转换为 .docx（保留标题层级、表格、列表、加粗、行内代码等）。"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else "/Users/duchengchao/work/demo/sxkj/Pet AI/需求设计文档/数心智能宠物健康平台_产品需求规格说明书_宠物主端与平台运营端_V1.0.md"
OUT = sys.argv[2] if len(sys.argv) > 2 else "/Users/duchengchao/work/demo/sxkj/Pet AI/需求设计文档/数心智能宠物健康平台_产品需求规格说明书_宠物主端与平台运营端_V1.0.docx"

BODY_FONT = "微软雅黑"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # 深蓝，标题色
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)

def set_run_font(run, font=BODY_FONT, size=10.5, bold=False, color=BODY_COLOR, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def add_inline(paragraph, text, base_size=10.5, base_bold=False, base_color=BODY_COLOR):
    """解析行内 **加粗** 与 `行内代码`。"""
    # 先按行内代码切分
    token_re = re.compile(r"(\*\*.*?\*\*|`[^`]*`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run_font(r, size=base_size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            set_run_font(r, font=MONO_FONT, size=base_size - 0.5, color=RGBColor(0xC0, 0x39, 0x2B))
        else:
            r = paragraph.add_run(part)
            set_run_font(r, size=base_size, bold=base_bold, color=base_color)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    add_inline(p, text, base_size=size, base_bold=bold, base_color=color or BODY_COLOR)

def add_heading(doc, text, level):
    sizes = {1: 20, 2: 16, 3: 13.5, 4: 12, 5: 11}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=ACCENT if level <= 3 else RGBColor(0x22, 0x22, 0x22))
    # 标题用底部边框增强层级感（1/2 级）
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1F4E79")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text)

def add_table(doc, header, rows):
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        set_cell_text(hdr[i], h.strip(), bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    # 数据行（斑马纹）
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, cell in enumerate(row):
            set_cell_text(cells[ci], cell.strip(), size=10)
            if ri % 2 == 1:
                shade_cell(cells[ci], "F2F6FB")
    # 行头加粗：若某列内容以「功能描述」等字段名结尾的列不处理，保持简单
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # 页面边距
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)

    # 默认正文样式
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结尾 ```
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(cl if cl else " ")
                set_run_font(r, font=MONO_FONT, size=8.5, color=RGBColor(0x33, 0x33, 0x33))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            # 收集表格块
            block = []
            j = i
            while j < n and lines[j].strip().startswith("|"):
                block.append(lines[j].strip())
                j += 1
            # 解析
            def split_row(r):
                r = r.strip()
                if r.startswith("|"):
                    r = r[1:]
                if r.endswith("|"):
                    r = r[:-1]
                return [c.strip() for c in r.split("|")]
            header = split_row(block[0])
            rows = [split_row(r) for r in block[2:]]  # 跳过分隔行
            add_table(doc, header, rows)
            i = j
            continue

        # 水平线
        if re.match(r"^---+$", stripped):
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_inline(p, text)
            i += 1
            continue

        # 无序列表
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            add_bullet(doc, m.group(1).strip())
            i += 1
            continue

        # 有序列表
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, m.group(1).strip())
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("OK ->", OUT)

if __name__ == "__main__":
    main()
